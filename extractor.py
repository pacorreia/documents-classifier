"""Text extraction from PDF and DOCX files."""

import logging
import os
import sys
from pathlib import Path

import pdfplumber
from docx import Document

# pdfminer (used by pdfplumber) logs harmless warnings such as:
#   "Could not get FontBBox from font descriptor because None cannot be
#    parsed as 4 floats"
# when a PDF font descriptor omits or corrupts the FontBBox entry.  The
# warning is cosmetic — text extraction is unaffected — so we suppress
# everything below ERROR from pdfminer.
logging.getLogger("pdfminer").setLevel(logging.ERROR)

_OCR_CHAR_THRESHOLD = 50  # pages with fewer chars trigger OCR


def _ocr_pdf(path: Path) -> str:
    """Rasterise every page and run Tesseract OCR on it."""
    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        return ""

    poppler_path: str | None = None

    # When running as a PyInstaller --onefile bundle, the bundled binaries
    # are extracted at runtime to sys._MEIPASS (a temp dir), NOT next to
    # sys.executable.  Using sys.executable.parent would point at the
    # binary itself and never find tesseract or pdftoppm.
    if getattr(sys, 'frozen', False):
        _base = Path(getattr(sys, '_MEIPASS', Path(sys.executable).parent))
        pytesseract.pytesseract.tesseract_cmd = str(_base / 'tesseract')
        os.environ['TESSDATA_PREFIX'] = str(_base / 'tessdata')
        poppler_path = str(_base)

    try:
        images = convert_from_path(path, dpi=300, poppler_path=poppler_path)
        return "\n".join(pytesseract.image_to_string(img) for img in images)
    except Exception:
        return ""


def extract_text_pdf(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text() or ""
            text_parts.append(extracted)

    full_text = "\n".join(text_parts)
    if len(full_text.strip()) < _OCR_CHAR_THRESHOLD:
        # Likely a scanned / image-based PDF — fall back to OCR.
        full_text = _ocr_pdf(path)
    return full_text


def _cell_text_recursive(cell, seen_tcs: set) -> str:
    """Return text from a cell and any nested tables within it.

    Uses the actual lxml _tc element (not id()) for deduplication so that
    garbage-collected proxy objects don't cause false duplicate hits.
    """
    parts = []
    if cell.text.strip():
        parts.append(cell.text)
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                tc = nested_cell._tc
                if tc not in seen_tcs:
                    seen_tcs.add(tc)
                    text = _cell_text_recursive(nested_cell, seen_tcs)
                    if text:
                        parts.append(text)
    return "\n".join(parts)


def extract_text_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # Store the actual lxml _tc element (strong reference) to prevent GC from
    # reusing memory addresses and causing false duplicate hits across cells.
    seen_tcs: set = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if tc not in seen_tcs:
                    seen_tcs.add(tc)
                    text = _cell_text_recursive(cell, seen_tcs)
                    if text:
                        parts.append(text)
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_pdf(path)
    if suffix == ".docx":
        return extract_text_docx(path)
    if suffix == ".doc":
        raise ValueError(
            f"{path.name}: legacy .doc format is not supported "
            "— please convert to .docx first."
        )
    raise ValueError(f"Unsupported file type: {suffix}")
